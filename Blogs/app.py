from os import error
from flask import Flask, render_template, Response, request, redirect
from includes import db

#from hackathon_3.includes.db import get_all
app = Flask(__name__)

# Jinja
@app.route("/")
def index():
    error = False
    blogs = db.get_all()
    return render_template("index.html", blogs=blogs, error = error)

@app.route('/', methods=["POST"])
def add():
    error = False
    title = request.form["title"]
    content = request.form["content"]   
    if title and content:
        db.insert(title, content)
    else:
        error = True 
    blogs = db.get_all()   
    return render_template("index.html", blogs=blogs, error = error)


@app.route("/about")
def about():
    return render_template("about.html")

@app.route("/resignation-letter")
def resignation_letter():
    error = False
    return render_template("resignation_letter.html", file_name="", error= error)    

@app.route("/resignation-letter", methods=["POST"])
def resignation_letter_export():
    from pathlib import Path
    
    from docx import Document
    from docx.shared import Cm
    fullname = request.form["fullname"]
    reason = request.form["reason"]
    file_name = ""
    error = False
    if fullname and reason:
        doc = Document()
        heading = doc.add_heading('Resignation Letter', level=1)
        doc.add_paragraph(fullname)
        doc.add_paragraph(reason)
        file_name = 'resignation-letter.docx'
        doc.save(str(Path(__file__).parent.absolute())+'/static/resignation-letter.docx')
    else:
        error= True
    return render_template("resignation_letter.html", file_name=file_name, error= error) 

# path param
@app.route("/post/<id>")
def post(id):
    blog = db.get_one(id)
    if blog: status_code = 200
    else: status_code = 404
    error = False
    return Response(render_template("edit.html", blog=blog, error = error), status=status_code, mimetype="text/html")


@app.route('/post/<id>', methods=["POST"])
def update(id):
    title = request.form["title"]
    content = request.form["content"]
    error = False
    if title and content:
        db.update(id, title, content)
        return redirect('/', 302)
    else:
        error= True     
        blog = db.get_one(id)
        if blog: status_code = 200
        else: status_code = 404
        return Response(render_template("edit.html", blog=blog, error = error), status=status_code, mimetype="text/html")



if __name__ == "__main__":
    app.run(debug=True)